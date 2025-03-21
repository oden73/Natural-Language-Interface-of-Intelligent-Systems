from docx import Document

from controller.format_controller import FormatController


class DOCXController(FormatController):
    def __init__(self) -> None:
        super().__init__()

    def file_parse(self, file_name: str) -> str:
        if not super()._is_file_valid(file_name):
            return ""

        doc: Document = Document(file_name)

        paragraphs = [par.text for par in doc.paragraphs]
        print(f'Найдено параграфов: {len(paragraphs)}')

        text = ''
        for par in doc.paragraphs:
            par_text: str = par.text
            if par_text:
                text += par_text + ' '

        text = super()._clean_text(text)
        return text

